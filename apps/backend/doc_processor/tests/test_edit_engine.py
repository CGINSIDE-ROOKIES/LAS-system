from __future__ import annotations

from io import BytesIO
from pathlib import Path
import tempfile
import unittest
from unittest.mock import patch
import zipfile

from document_processor import DocIR

from doc_processor.edit_engine import (
    EditValidationError,
    ParagraphTextEdit,
    RunTextEdit,
    apply_edits_to_doc_ir,
    apply_edits_to_file,
    validate_edit_commands,
)


class EditEngineTests(unittest.TestCase):
    @staticmethod
    def _ids(doc: DocIR) -> dict[str, str]:
        return {
            "p1": doc.paragraphs[0].node_id,
            "r1": doc.paragraphs[0].runs[0].node_id,
            "r2": doc.paragraphs[0].runs[1].node_id if len(doc.paragraphs[0].runs) > 1 else "",
        }

    @staticmethod
    def _build_simple_hwpx_bytes() -> bytes:
        hwpx_bytes = BytesIO()
        with zipfile.ZipFile(hwpx_bytes, "w") as archive:
            archive.writestr(
                "Contents/header.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" />
""",
            )
            archive.writestr(
                "Contents/section0.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p>
    <hp:run><hp:t>Hello </hp:t></hp:run>
    <hp:run><hp:t>World</hp:t></hp:run>
  </hp:p>
</hs:sec>
""",
            )
        return hwpx_bytes.getvalue()

    @staticmethod
    def _build_single_run_hwpx_bytes(text: str) -> bytes:
        hwpx_bytes = BytesIO()
        with zipfile.ZipFile(hwpx_bytes, "w") as archive:
            archive.writestr(
                "Contents/header.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" />
""",
            )
            archive.writestr(
                "Contents/section0.xml",
                f"""<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p>
    <hp:run><hp:t>{text}</hp:t></hp:run>
  </hp:p>
</hs:sec>
""",
            )
        return hwpx_bytes.getvalue()

    def test_apply_edits_to_doc_ir_updates_text_and_reports_modified_runs(self) -> None:
        doc = DocIR.from_mapping(
            {
                "s1.p1.r1": "Hello ",
                "s1.p1.r2": "World",
            },
            source_doc_type="docx",
        )
        ids = self._ids(doc)

        updated, result = apply_edits_to_doc_ir(
            doc,
            [
                ParagraphTextEdit(
                    paragraph_id=ids["p1"],
                    old_text="Hello World",
                    new_text="Hello Legal World",
                    reason="Insert qualifier",
                )
            ],
        )

        self.assertEqual(updated.paragraphs[0].text, "Hello Legal World")
        self.assertEqual(result.edits_applied, 1)
        self.assertIn(ids["p1"], result.modified_target_ids)
        self.assertIn(ids["r1"], result.modified_run_ids)

    def test_apply_edits_to_doc_ir_replaces_single_run_paragraph_without_garbling(self) -> None:
        doc = DocIR.from_mapping(
            {
                "s1.p1.r1": "제3조 (계약기간 등)",
            },
            source_doc_type="hwpx",
        )
        ids = self._ids(doc)

        updated, result = apply_edits_to_doc_ir(
            doc,
            [
                ParagraphTextEdit(
                    paragraph_id=ids["p1"],
                    old_text="제3조 (계약기간 등)",
                    new_text="updated paragraph text",
                    reason="Replace heading",
                )
            ],
        )

        self.assertEqual(updated.paragraphs[0].text, "updated paragraph text")
        self.assertEqual(updated.paragraphs[0].runs[0].text, "updated paragraph text")
        self.assertEqual(result.edits_applied, 1)
        self.assertIn(ids["r1"], result.modified_run_ids)

    def test_validate_edit_commands_rejects_text_mismatch(self) -> None:
        doc = DocIR.from_mapping({"s1.p1.r1": "Hello"})
        ids = self._ids(doc)

        with self.assertRaises(EditValidationError):
            validate_edit_commands(
                doc,
                [RunTextEdit(run_id=ids["r1"], old_text="World", new_text="Hello")],
            )

    def test_apply_edits_to_file_updates_docx_and_preserves_run_style(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "sample.docx"
            output = Path(tmp_dir) / "sample_edited.docx"

            docx = Document()
            paragraph = docx.add_paragraph()
            paragraph.add_run("Hello ")
            styled = paragraph.add_run("World")
            styled.bold = True
            docx.save(source)
            ids = self._ids(DocIR.from_file(source))

            result = apply_edits_to_file(
                source,
                [
                    ParagraphTextEdit(
                        paragraph_id=ids["p1"],
                        old_text="Hello World",
                        new_text="Hello Legal World",
                        reason="Expand wording",
                    )
                ],
                output_path=output,
            )

            edited_docx = Document(output)
            self.assertEqual(result.output_path, str(output))
            self.assertEqual(edited_docx.paragraphs[0].text, "Hello Legal World")
            self.assertTrue(edited_docx.paragraphs[0].runs[1].bold)

            reparsed = DocIR.from_file(output)
            self.assertEqual(reparsed.paragraphs[0].text, "Hello Legal World")

    def test_apply_edits_to_file_updates_hwpx_archive(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "sample.hwpx"
            output = Path(tmp_dir) / "sample_edited.hwpx"
            source.write_bytes(self._build_simple_hwpx_bytes())
            ids = self._ids(DocIR.from_file(source))

            result = apply_edits_to_file(
                source,
                [
                    RunTextEdit(
                        run_id=ids["r2"],
                        old_text="World",
                        new_text="HWPX",
                        reason="Rename token",
                    )
                ],
                output_path=output,
            )

            self.assertEqual(result.output_path, str(output))
            reparsed = DocIR.from_file(output)
            self.assertEqual(reparsed.paragraphs[0].text, "Hello HWPX")

    def test_apply_edits_to_file_refuses_to_overwrite_source(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "sample.docx"

            docx = Document()
            docx.add_paragraph("Hello World")
            docx.save(source)
            ids = self._ids(DocIR.from_file(source))

            with self.assertRaises(EditValidationError):
                apply_edits_to_file(
                    source,
                    [
                        ParagraphTextEdit(
                            paragraph_id=ids["p1"],
                            old_text="Hello World",
                            new_text="Hello Contract World",
                            reason="expand wording",
                        )
                    ],
                    output_path=source,
                )

    def test_apply_edits_to_file_replaces_single_run_hwpx_paragraph_without_garbling(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "sample.hwpx"
            output = Path(tmp_dir) / "sample_edited.hwpx"
            source.write_bytes(self._build_single_run_hwpx_bytes("제3조 (계약기간 등)"))
            ids = self._ids(DocIR.from_file(source))

            result = apply_edits_to_file(
                source,
                [
                    ParagraphTextEdit(
                        paragraph_id=ids["p1"],
                        old_text="제3조 (계약기간 등)",
                        new_text="updated paragraph text",
                        reason="Replace heading",
                    )
                ],
                output_path=output,
            )

            self.assertEqual(result.output_path, str(output))
            reparsed = DocIR.from_file(output)
            self.assertEqual(reparsed.paragraphs[0].text, "updated paragraph text")
            self.assertEqual(reparsed.paragraphs[0].runs[0].text, "updated paragraph text")

    def test_apply_edits_to_file_converts_hwp_source_and_writes_hwpx(self) -> None:
        fake_doc = DocIR.from_mapping(
            {
                "s1.p1.r1": "Hello ",
                "s1.p1.r2": "World",
            },
            source_doc_type="hwp",
        )
        ids = self._ids(fake_doc)

        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "sample.hwp"
            output = Path(tmp_dir) / "sample_edited.hwp"
            source.write_bytes(b"fake-hwp")

            with (
                patch("document_processor.edit_engine.DocIR.from_file", return_value=fake_doc),
                patch(
                    "document_processor.edit_engine.convert_hwp_to_hwpx_bytes",
                    return_value=self._build_simple_hwpx_bytes(),
                ) as convert_hwp,
            ):
                result = apply_edits_to_file(
                    source,
                    [
                        RunTextEdit(
                            run_id=ids["r2"],
                            old_text="World",
                            new_text="HWPX",
                            reason="Rename token",
                        )
                    ],
                    output_path=output,
                )

            expected_output = output.with_suffix(".hwpx")
            self.assertEqual(result.output_path, str(expected_output))
            self.assertTrue(result.warnings)
            self.assertIn("HWP sources are written back as HWPX", result.warnings[0])

            convert_source = convert_hwp.call_args.kwargs.get("hwp_path", convert_hwp.call_args.args[0])
            self.assertEqual(Path(convert_source), source)

            reparsed = DocIR.from_file(expected_output)
            self.assertEqual(reparsed.paragraphs[0].text, "Hello HWPX")


if __name__ == "__main__":
    unittest.main()
