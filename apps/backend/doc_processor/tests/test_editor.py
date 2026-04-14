from __future__ import annotations

from io import BytesIO
from pathlib import Path
import tempfile
import unittest
import zipfile

from document_processor import DocIR

from doc_processor.editor import (
    EditValidationError,
    ParagraphTextEdit,
    RunTextEdit,
    apply_edits_to_doc_ir,
    apply_edits_to_file,
    validate_edit_commands,
)


class EditorTests(unittest.TestCase):
    def test_apply_edits_to_doc_ir_updates_text_and_reports_modified_runs(self) -> None:
        doc = DocIR.from_mapping(
            {
                "s1.p1.r1": "Hello ",
                "s1.p1.r2": "World",
            },
            source_doc_type="docx",
        )

        updated, result = apply_edits_to_doc_ir(
            doc,
            [
                ParagraphTextEdit(
                    paragraph_unit_id="s1.p1",
                    old_text="Hello World",
                    new_text="Hello Legal World",
                    reason="Insert qualifier",
                )
            ],
        )

        self.assertEqual(updated.paragraphs[0].text, "Hello Legal World")
        self.assertEqual(result.edits_applied, 1)
        self.assertIn("s1.p1", result.modified_unit_ids)
        self.assertIn("s1.p1.r1", result.modified_run_ids)

    def test_validate_edit_commands_rejects_text_mismatch(self) -> None:
        doc = DocIR.from_mapping({"s1.p1.r1": "Hello"})

        with self.assertRaises(EditValidationError):
            validate_edit_commands(
                doc,
                [RunTextEdit(run_unit_id="s1.p1.r1", old_text="World", new_text="Hello")],
            )

    def test_apply_edits_to_file_updates_docx_and_preserves_run_style(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "sample.docx"
            output = Path(tmp_dir) / "sample_edited.docx"

            docx = Document()
            paragraph = docx.add_paragraph()
            paragraph.add_run("Hello ")
            styled = paragraph.add_run("World")
            styled.bold = True
            docx.save(source)

            result = apply_edits_to_file(
                source,
                [
                    ParagraphTextEdit(
                        paragraph_unit_id="s1.p1",
                        old_text="Hello World",
                        new_text="Hello Legal World",
                        reason="Expand wording",
                    )
                ],
                output_path=output,
            )

            edited_docx = Document(output)
            self.assertEqual(result.output_path, str(output))
            self.assertEqual(edited_docx.paragraphs[0].text, "Hello Legal World")
            self.assertTrue(edited_docx.paragraphs[0].runs[1].bold)

            reparsed = DocIR.from_file(output)
            self.assertEqual(reparsed.paragraphs[0].text, "Hello Legal World")

    def test_apply_edits_to_file_updates_hwpx_archive(self) -> None:
        hwpx_bytes = BytesIO()
        with zipfile.ZipFile(hwpx_bytes, "w") as archive:
            archive.writestr(
                "Contents/header.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" />
""",
            )
            archive.writestr(
                "Contents/section0.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p>
    <hp:run><hp:t>Hello </hp:t></hp:run>
    <hp:run><hp:t>World</hp:t></hp:run>
  </hp:p>
</hs:sec>
""",
            )

        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "sample.hwpx"
            output = Path(tmp_dir) / "sample_edited.hwpx"
            source.write_bytes(hwpx_bytes.getvalue())

            result = apply_edits_to_file(
                source,
                [
                    RunTextEdit(
                        run_unit_id="s1.p1.r2",
                        old_text="World",
                        new_text="HWPX",
                        reason="Rename token",
                    )
                ],
                output_path=output,
            )

            self.assertEqual(result.output_path, str(output))
            reparsed = DocIR.from_file(output)
            self.assertEqual(reparsed.paragraphs[0].text, "Hello HWPX")


if __name__ == "__main__":
    unittest.main()
