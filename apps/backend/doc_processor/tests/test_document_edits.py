from __future__ import annotations

from io import BytesIO
import hashlib
from pathlib import Path
import tempfile
import unittest
import zipfile

from document_processor import DocIR

from doc_processor.api import DocumentInput, TextEdit, apply_document_edits, validate_document_edits


class DocumentEditTests(unittest.TestCase):
    @staticmethod
    def _text_hash(text: str) -> str:
        return hashlib.sha1(text.encode("utf-8")).hexdigest()

    @staticmethod
    def _ids(doc: DocIR) -> dict[str, str]:
        return {
            "p1": doc.paragraphs[0].node_id,
            "r1": doc.paragraphs[0].runs[0].node_id,
            "r2": doc.paragraphs[0].runs[1].node_id if len(doc.paragraphs[0].runs) > 1 else "",
        }

    @staticmethod
    def _build_simple_hwpx_bytes() -> bytes:
        hwpx_bytes = BytesIO()
        with zipfile.ZipFile(hwpx_bytes, "w") as archive:
            archive.writestr(
                "Contents/header.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" />
""",
            )
            archive.writestr(
                "Contents/section0.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p>
    <hp:run><hp:t>Hello </hp:t></hp:run>
    <hp:run><hp:t>World</hp:t></hp:run>
  </hp:p>
</hs:sec>
""",
            )
        return hwpx_bytes.getvalue()

    def test_apply_document_edits_updates_doc_ir_and_reports_modified_runs(self) -> None:
        doc = DocIR.from_mapping(
            {
                "s1.p1.r1": "Hello ",
                "s1.p1.r2": "World",
            },
            source_doc_type="docx",
        )
        ids = self._ids(doc)

        result = apply_document_edits(
            document=DocumentInput(doc_ir=doc),
            edits=[
                TextEdit(
                    target_kind="paragraph",
                    target_id=ids["p1"],
                    expected_text_hash=self._text_hash("Hello World"),
                    new_text="Hello Legal World",
                    reason="Insert qualifier",
                )
            ],
            return_doc_ir=True,
        )

        self.assertTrue(result.ok)
        self.assertIsNotNone(result.updated_doc_ir)
        self.assertEqual(result.updated_doc_ir.paragraphs[0].text, "Hello Legal World")
        self.assertEqual(result.edits_applied, 1)
        self.assertIn(ids["p1"], result.modified_target_ids)
        self.assertIn(ids["r1"], result.modified_run_ids)

    def test_validate_document_edits_rejects_text_hash_mismatch(self) -> None:
        doc = DocIR.from_mapping({"s1.p1.r1": "Hello"})
        ids = self._ids(doc)

        result = validate_document_edits(
            document=DocumentInput(doc_ir=doc),
            edits=[
                TextEdit(
                    target_kind="run",
                    target_id=ids["r1"],
                    expected_text_hash=self._text_hash("World"),
                    new_text="Hello",
                )
            ],
        )

        self.assertFalse(result.ok)
        self.assertEqual(result.issues[0].code, "text_hash_mismatch")

    def test_apply_document_edits_updates_docx_and_preserves_run_style(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "sample.docx"
            output = Path(tmp_dir) / "sample_edited.docx"

            docx = Document()
            paragraph = docx.add_paragraph()
            paragraph.add_run("Hello ")
            styled = paragraph.add_run("World")
            styled.bold = True
            docx.save(source)
            ids = self._ids(DocIR.from_file(source))

            result = apply_document_edits(
                source_path=source,
                edits=[
                    TextEdit(
                        target_kind="paragraph",
                        target_id=ids["p1"],
                        expected_text_hash=self._text_hash("Hello World"),
                        new_text="Hello Legal World",
                        reason="Expand wording",
                    )
                ],
                output_path=str(output),
            )

            edited_docx = Document(output)
            self.assertTrue(result.ok)
            self.assertEqual(result.output_path, str(output))
            self.assertEqual(edited_docx.paragraphs[0].text, "Hello Legal World")
            self.assertTrue(edited_docx.paragraphs[0].runs[1].bold)

    def test_apply_document_edits_updates_hwpx_archive(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "sample.hwpx"
            output = Path(tmp_dir) / "sample_edited.hwpx"
            source.write_bytes(self._build_simple_hwpx_bytes())
            ids = self._ids(DocIR.from_file(source))

            result = apply_document_edits(
                source_path=source,
                edits=[
                    TextEdit(
                        target_kind="run",
                        target_id=ids["r2"],
                        expected_text_hash=self._text_hash("World"),
                        new_text="HWPX",
                        reason="Rename token",
                    )
                ],
                output_path=str(output),
            )

            self.assertTrue(result.ok)
            self.assertEqual(result.output_path, str(output))
            reparsed = DocIR.from_file(output)
            self.assertEqual(reparsed.paragraphs[0].text, "Hello HWPX")


if __name__ == "__main__":
    unittest.main()
