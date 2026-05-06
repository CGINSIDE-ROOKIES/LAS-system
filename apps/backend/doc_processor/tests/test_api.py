from __future__ import annotations

import tempfile
import unittest
import hashlib
from pathlib import Path

from document_processor import DocIR

from doc_processor import DocumentInput, RelevanceMode
from doc_processor.api import (
    TextAnnotation,
    TextEdit,
    apply_document_edits,
    get_document_context,
    list_editable_targets,
    parse_document,
    render_review_html,
    validate_document_edits,
)


ROOT = Path(__file__).resolve().parents[1]
DOC_SAMPLES = ROOT / "tests" / "doc_samples" / "new_test"


class ApiTests(unittest.TestCase):
    @staticmethod
    def _text_hash(text: str) -> str:
        return hashlib.sha1(text.encode("utf-8")).hexdigest()

    @staticmethod
    def _build_sample_docx(path: Path) -> None:
        from docx import Document

        docx = Document()
        paragraph = docx.add_paragraph()
        paragraph.add_run("Hello ")
        paragraph.add_run("World")
        docx.add_paragraph("Second paragraph")
        docx.save(path)

    @staticmethod
    def _sample_ids(path: Path) -> dict[str, str]:
        doc = DocIR.from_file(path)
        return {
            "p1": doc.paragraphs[0].node_id,
            "p2": doc.paragraphs[1].node_id,
            "r1": doc.paragraphs[0].runs[0].node_id,
            "r2": doc.paragraphs[0].runs[1].node_id,
        }

    def test_parse_document_returns_compact_llm_friendly_summary(self) -> None:
        target = DOC_SAMPLES / "02. 청소년 대중문화예술인 표준 부속합의서.hwpx"
        result = parse_document(
            source_path=str(target),
            relevance_mode=RelevanceMode.DISABLED,
            boundary_review_enabled=False,
            label_review_enabled=False,
            include_editable_targets=True,
            max_editable_targets=5,
        )

        self.assertTrue(result.accepted)
        self.assertGreater(len(result.paragraphs), 0)
        self.assertGreater(len(result.clauses), 0)
        self.assertGreater(len(result.editable_targets), 0)

    def test_parse_document_accepts_bytes_backed_input(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "sample.docx"
            self._build_sample_docx(source)

            result = parse_document(
                document=DocumentInput(
                    source_bytes=source.read_bytes(),
                    source_name=source.name,
                    source_doc_type="docx",
                ),
                relevance_mode=RelevanceMode.DISABLED,
                boundary_review_enabled=False,
                label_review_enabled=False,
                include_editable_targets=True,
                max_editable_targets=5,
            )

        self.assertTrue(result.accepted)
        self.assertEqual(result.source_name, "sample.docx")
        self.assertGreater(len(result.paragraphs), 0)

    def test_get_document_context_and_editable_targets_use_exact_text(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "sample.docx"
            self._build_sample_docx(source)
            ids = self._sample_ids(source)

            context = get_document_context(
                source_path=str(source),
                target_ids=[ids["r2"]],
                before=0,
                after=1,
            )
            self.assertEqual([paragraph.node_id for paragraph in context.paragraphs], [ids["p1"], ids["p2"]])
            self.assertEqual(context.paragraphs[0].runs[1].text, "World")

            targets = list_editable_targets(
                source_path=str(source),
                target_ids=[ids["p1"]],
                target_kinds=["run"],
                include_child_runs=True,
            )
            self.assertEqual([target.target_id for target in targets.targets], [ids["r1"], ids["r2"]])

    def test_validate_document_edits_returns_text_hash_issue(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "sample.docx"
            self._build_sample_docx(source)
            ids = self._sample_ids(source)

            result = validate_document_edits(
                source_path=str(source),
                edits=[
                    TextEdit(
                        target_kind="paragraph",
                        target_id=ids["p1"],
                        expected_text_hash=self._text_hash("World"),
                        new_text="Hello Legal World",
                        reason="wrong old text",
                    )
                ],
            )

            self.assertFalse(result.ok)
            self.assertEqual([issue.code for issue in result.issues], ["text_hash_mismatch"])
            self.assertEqual(result.issues[0].current_text, "Hello World")

    def test_validate_document_edits_returns_target_kind_issue(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "sample.docx"
            self._build_sample_docx(source)
            ids = self._sample_ids(source)

            result = validate_document_edits(
                source_path=str(source),
                edits=[
                    TextEdit(
                        target_kind="run",
                        target_id=ids["p1"],
                        expected_text_hash=self._text_hash("Hello World"),
                        new_text="Hello Legal World",
                        reason="wrong target kind",
                    )
                ],
            )

            self.assertFalse(result.ok)
            self.assertEqual([issue.code for issue in result.issues], ["target_kind_mismatch"])

    def test_apply_document_edits_applies_and_reports_modified_targets(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "sample.docx"
            output = Path(tmp_dir) / "sample_edited.docx"
            self._build_sample_docx(source)
            ids = self._sample_ids(source)

            result = apply_document_edits(
                source_path=str(source),
                output_path=str(output),
                edits=[
                    TextEdit(
                        target_kind="paragraph",
                        target_id=ids["p1"],
                        expected_text_hash=self._text_hash("Hello World"),
                        new_text="Hello Legal World",
                        reason="expand wording",
                    )
                ],
            )

            self.assertTrue(result.ok)
            self.assertEqual(result.output_path, str(output))
            self.assertIn(ids["p1"], result.modified_target_ids)

            reparsed = DocIR.from_file(output)
            self.assertEqual(reparsed.paragraphs[0].text, "Hello Legal World")

    def test_apply_document_edits_supports_output_filename_next_to_source(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "sample.docx"
            output = Path(tmp_dir) / "llm_review.docx"
            self._build_sample_docx(source)
            ids = self._sample_ids(source)

            result = apply_document_edits(
                source_path=str(source),
                output_filename="llm_review.docx",
                edits=[
                    TextEdit(
                        target_kind="paragraph",
                        target_id=ids["p1"],
                        expected_text_hash=self._text_hash("Hello World"),
                        new_text="Hello Contract World",
                        reason="expand wording",
                    )
                ],
            )

            self.assertTrue(result.ok)
            self.assertEqual(result.output_path, str(output))
            self.assertEqual(DocIR.from_file(output).paragraphs[0].text, "Hello Contract World")

    def test_apply_document_edits_rejects_output_path_that_matches_source(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "sample.docx"
            self._build_sample_docx(source)
            ids = self._sample_ids(source)

            result = apply_document_edits(
                source_path=str(source),
                output_path=str(source),
                edits=[
                    TextEdit(
                        target_kind="paragraph",
                        target_id=ids["p1"],
                        expected_text_hash=self._text_hash("Hello World"),
                        new_text="Hello Contract World",
                        reason="expand wording",
                    )
                ],
            )

            self.assertFalse(result.ok)
            self.assertEqual(result.validation.issues[0].code, "output_path_conflicts_with_source")
            self.assertEqual(DocIR.from_file(source).paragraphs[0].text, "Hello World")

    def test_render_review_html_returns_structured_annotation_validation(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "sample.docx"
            self._build_sample_docx(source)
            ids = self._sample_ids(source)

            ok_result = render_review_html(
                source_path=str(source),
                annotations=[
                    TextAnnotation(
                        target_kind="paragraph",
                        target_id=ids["p1"],
                        selected_text="Hello",
                        label="Intro",
                    )
                ],
            )
            self.assertTrue(ok_result.ok)
            self.assertIn("<mark", ok_result.html or "")
            self.assertEqual(ok_result.resolved_annotations[0].selected_text, "Hello")

            bad_result = render_review_html(
                source_path=str(source),
                annotations=[
                    TextAnnotation(
                        target_kind="run",
                        target_id=ids["p1"],
                        selected_text="Hello",
                        label="Wrong kind",
                    )
                ],
            )
            self.assertFalse(bad_result.ok)
            self.assertEqual(bad_result.validation.issues[0].code, "target_kind_mismatch")

    def test_render_review_html_returns_ambiguity_validation_for_repeated_text(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "sample.docx"
            self._build_sample_docx(source)
            ids = self._sample_ids(source)

            result = render_review_html(
                source_path=str(source),
                annotations=[
                    TextAnnotation(
                        target_kind="paragraph",
                        target_id=ids["p1"],
                        selected_text="l",
                        label="Ambiguous letter",
                    )
                ],
            )

            self.assertFalse(result.ok)
            self.assertEqual(result.validation.issues[0].code, "selected_text_ambiguous")


if __name__ == "__main__":
    unittest.main()
