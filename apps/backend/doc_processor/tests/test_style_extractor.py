from __future__ import annotations

from io import BytesIO
import sys
import tempfile
import unittest
from pathlib import Path
import zipfile

THIS_DIR = Path(__file__).resolve().parent
PACKAGE_ROOT = THIS_DIR.parent
if str(PACKAGE_ROOT) not in sys.path:
    sys.path.insert(0, str(PACKAGE_ROOT))

from core.style_extractor import extract_styles, extract_styles_docx, extract_styles_hwpx


class StyleExtractorTests(unittest.TestCase):
    def test_extract_docx_paragraph_indents(self) -> None:
        from docx import Document
        from docx.shared import Pt

        with tempfile.TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "indent_sample.docx"

            doc = Document()
            paragraph = doc.add_paragraph("Indented")
            paragraph.alignment = 1  # center
            paragraph.paragraph_format.left_indent = Pt(24.0)
            paragraph.paragraph_format.right_indent = Pt(12.0)
            paragraph.paragraph_format.first_line_indent = Pt(-6.0)
            doc.save(str(docx_path))

            style_map = extract_styles_docx(docx_path)

        pstyle = style_map.paragraphs.get("s1.p1")
        self.assertIsNotNone(pstyle)
        assert pstyle is not None

        self.assertEqual(pstyle.align, "center")
        self.assertAlmostEqual(pstyle.left_indent_pt or 0.0, 24.0, places=3)
        self.assertAlmostEqual(pstyle.right_indent_pt or 0.0, 12.0, places=3)
        self.assertAlmostEqual(pstyle.first_line_indent_pt or 0.0, -6.0, places=3)
        self.assertAlmostEqual(pstyle.hanging_indent_pt or 0.0, 6.0, places=3)

    def test_extract_hwpx_paragraph_indents(self) -> None:
        header_xml = """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core">
  <hh:paraProperties itemCnt="1">
    <hh:paraPr id="1">
      <hh:align horizontal="CENTER" />
      <hh:margin>
        <hc:intent value="-500" unit="HWPUNIT" />
        <hc:left value="200" unit="HWPUNIT" />
        <hc:right value="300" unit="HWPUNIT" />
      </hh:margin>
    </hh:paraPr>
  </hh:paraProperties>
  <hh:charProperties itemCnt="1">
    <hh:charPr id="1" height="1200" textColor="#112233">
      <hh:bold />
      <hh:underline type="BOTTOM" shape="SOLID" color="#000000" />
      <hh:strikeout shape="NONE" color="#000000" />
    </hh:charPr>
  </hh:charProperties>
</hh:head>
"""

        section_xml = """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec
  xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section"
  xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph"
  xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head"
  xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core">
  <hp:p paraPrIDRef="1">
    <hp:run charPrIDRef="1">
      <hp:t>Hello</hp:t>
    </hp:run>
  </hp:p>
</hs:sec>
"""

        hwpx_bytes_io = BytesIO()
        with zipfile.ZipFile(hwpx_bytes_io, "w") as zf:
            zf.writestr("Contents/header.xml", header_xml)
            zf.writestr("Contents/section0.xml", section_xml)
        hwpx_bytes = hwpx_bytes_io.getvalue()

        style_map = extract_styles_hwpx(hwpx_bytes)

        pstyle = style_map.paragraphs.get("s1.p1")
        self.assertIsNotNone(pstyle)
        assert pstyle is not None

        self.assertEqual(pstyle.align, "center")
        self.assertAlmostEqual(pstyle.left_indent_pt or 0.0, 2.0, places=3)
        self.assertAlmostEqual(pstyle.right_indent_pt or 0.0, 3.0, places=3)
        self.assertAlmostEqual(pstyle.first_line_indent_pt or 0.0, -5.0, places=3)
        self.assertAlmostEqual(pstyle.hanging_indent_pt or 0.0, 5.0, places=3)

        rstyle = style_map.runs.get("s1.p1.r1")
        self.assertIsNotNone(rstyle)
        assert rstyle is not None
        self.assertTrue(rstyle.bold)
        self.assertTrue(rstyle.underline)
        self.assertEqual(rstyle.color, "#112233")
        self.assertAlmostEqual(rstyle.size_pt or 0.0, 12.0, places=3)

    def test_unified_interface_accepts_bytes_hwpx(self) -> None:
        hwpx_bytes_io = BytesIO()
        with zipfile.ZipFile(hwpx_bytes_io, "w") as zf:
            zf.writestr(
                "Contents/header.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core">
  <hh:paraProperties itemCnt="1">
    <hh:paraPr id="1">
      <hh:margin>
        <hc:intent value="100" unit="HWPUNIT" />
      </hh:margin>
    </hh:paraPr>
  </hh:paraProperties>
</hh:head>
""",
            )
            zf.writestr(
                "Contents/section0.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p paraPrIDRef="1"><hp:run><hp:t>X</hp:t></hp:run></hp:p>
</hs:sec>
""",
            )

        style_map = extract_styles(hwpx_bytes_io.getvalue())
        pstyle = style_map.paragraphs.get("s1.p1")
        self.assertIsNotNone(pstyle)
        assert pstyle is not None
        self.assertAlmostEqual(pstyle.first_line_indent_pt or 0.0, 1.0, places=3)


if __name__ == "__main__":
    unittest.main()
