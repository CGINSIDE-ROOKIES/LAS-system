from __future__ import annotations

import difflib
from io import BytesIO
from pathlib import Path
import re
from typing import Callable
from xml.etree import ElementTree as ET
import zipfile

from pydantic import BaseModel, Field

from document_processor import DocIR, ParagraphIR, RunIR, TableIR


class EditValidationError(ValueError):
    pass


class RunTextEdit(BaseModel):
    run_unit_id: str
    old_text: str
    new_text: str
    reason: str = ""


class ParagraphTextEdit(BaseModel):
    paragraph_unit_id: str
    old_text: str
    new_text: str
    reason: str = ""


EditCommand = RunTextEdit | ParagraphTextEdit


class ApplyEditsResult(BaseModel):
    source_doc_type: str | None = None
    output_path: str | None = None
    edits_applied: int = 0
    modified_unit_ids: list[str] = Field(default_factory=list)
    modified_run_ids: list[str] = Field(default_factory=list)
    warnings: list[str] = Field(default_factory=list)


class _EditableRunRef:
    def __init__(
        self,
        *,
        unit_id: str,
        get_text: Callable[[], str],
        set_text: Callable[[str], None],
    ) -> None:
        self.unit_id = unit_id
        self._get_text = get_text
        self._set_text = set_text

    @property
    def text(self) -> str:
        return self._get_text()

    @text.setter
    def text(self, value: str) -> None:
        self._set_text(value)


class _EditableParagraphRef:
    def __init__(
        self,
        *,
        unit_id: str,
        runs: list[_EditableRunRef],
        has_non_run_content: bool = False,
        recompute: Callable[[], None] | None = None,
    ) -> None:
        self.unit_id = unit_id
        self.runs = runs
        self.has_non_run_content = has_non_run_content
        self._recompute = recompute

    @property
    def text(self) -> str:
        return "".join(run.text for run in self.runs)

    def recompute(self) -> None:
        if self._recompute is not None:
            self._recompute()


class _EditableDocIndex:
    def __init__(
        self,
        *,
        paragraphs: dict[str, _EditableParagraphRef],
        runs: dict[str, _EditableRunRef],
        run_to_paragraph: dict[str, _EditableParagraphRef],
    ) -> None:
        self.paragraphs = paragraphs
        self.runs = runs
        self.run_to_paragraph = run_to_paragraph


class _RunSpan(BaseModel):
    start: int
    end: int
    run: _EditableRunRef

    model_config = {"arbitrary_types_allowed": True}


def _iter_doc_ir_paragraphs(paragraphs: list[ParagraphIR]):
    for paragraph in paragraphs:
        yield paragraph
        for table in paragraph.tables:
            yield from _iter_doc_ir_table_paragraphs(table)


def _iter_doc_ir_table_paragraphs(table: TableIR):
    for cell in table.cells:
        for paragraph in cell.paragraphs:
            yield paragraph
            for nested_table in paragraph.tables:
                yield from _iter_doc_ir_table_paragraphs(nested_table)


def _build_doc_ir_index(doc: DocIR) -> _EditableDocIndex:
    paragraphs: dict[str, _EditableParagraphRef] = {}
    runs: dict[str, _EditableRunRef] = {}
    run_to_paragraph: dict[str, _EditableParagraphRef] = {}

    for paragraph in _iter_doc_ir_paragraphs(doc.paragraphs):
        run_refs: list[_EditableRunRef] = []
        paragraph_ref = _EditableParagraphRef(
            unit_id=paragraph.unit_id,
            runs=run_refs,
            has_non_run_content=bool(paragraph.images or paragraph.tables),
            recompute=paragraph.recompute_text,
        )
        paragraphs[paragraph.unit_id] = paragraph_ref
        for run in paragraph.runs:
            run_ref = _EditableRunRef(
                unit_id=run.unit_id,
                get_text=lambda node=run: node.text,
                set_text=lambda value, node=run: setattr(node, "text", value),
            )
            run_refs.append(run_ref)
            runs[run.unit_id] = run_ref
            run_to_paragraph[run.unit_id] = paragraph_ref

    return _EditableDocIndex(paragraphs=paragraphs, runs=runs, run_to_paragraph=run_to_paragraph)


def _iter_docx_blocks(doc):
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    iter_inner_content = getattr(doc, "iter_inner_content", None)
    if callable(iter_inner_content):
        yield from iter_inner_content()
        return

    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def _iter_docx_blocks_from_element(parent, element):
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for child in element.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _build_docx_index(doc) -> _EditableDocIndex:
    paragraphs: dict[str, _EditableParagraphRef] = {}
    runs: dict[str, _EditableRunRef] = {}
    run_to_paragraph: dict[str, _EditableParagraphRef] = {}

    def register_paragraph(paragraph, paragraph_id: str, *, has_non_run_content: bool = False) -> _EditableParagraphRef:
        run_refs: list[_EditableRunRef] = []
        paragraph_ref = _EditableParagraphRef(
            unit_id=paragraph_id,
            runs=run_refs,
            has_non_run_content=has_non_run_content,
        )
        paragraphs[paragraph_id] = paragraph_ref
        for run_index, run in enumerate(paragraph.runs, start=1):
            run_id = f"{paragraph_id}.r{run_index}"
            run_ref = _EditableRunRef(
                unit_id=run_id,
                get_text=lambda node=run: node.text or "",
                set_text=lambda value, node=run: setattr(node, "text", value),
            )
            run_refs.append(run_ref)
            runs[run_id] = run_ref
            run_to_paragraph[run_id] = paragraph_ref
        return paragraph_ref

    def walk_table(table, table_base: str) -> None:
        for tr_idx, row in enumerate(table.rows, start=1):
            for tc_idx, cell in enumerate(row.cells, start=1):
                cp_idx = 0
                current_paragraph_id: str | None = None
                nested_table_counter_by_paragraph: dict[str, int] = {}

                for block in _iter_docx_blocks_from_element(cell, cell._tc):
                    if block.__class__.__name__ == "Paragraph":
                        cp_idx += 1
                        current_paragraph_id = f"{table_base}.tr{tr_idx}.tc{tc_idx}.p{cp_idx}"
                        register_paragraph(block, current_paragraph_id)
                        continue

                    if block.__class__.__name__ != "Table":
                        continue

                    if current_paragraph_id is None:
                        cp_idx += 1
                        current_paragraph_id = f"{table_base}.tr{tr_idx}.tc{tc_idx}.p{cp_idx}"
                        paragraph_ref = _EditableParagraphRef(
                            unit_id=current_paragraph_id,
                            runs=[],
                            has_non_run_content=True,
                        )
                        paragraphs[current_paragraph_id] = paragraph_ref
                    else:
                        paragraphs[current_paragraph_id].has_non_run_content = True

                    tbl_counter = nested_table_counter_by_paragraph.get(current_paragraph_id, 0) + 1
                    nested_table_counter_by_paragraph[current_paragraph_id] = tbl_counter
                    nested_table_base = f"{current_paragraph_id}.tbl{tbl_counter}"
                    walk_table(block, nested_table_base)

    p_idx = 0
    tbl_counter = 0
    for block in _iter_docx_blocks(doc):
        if block.__class__.__name__ == "Paragraph":
            p_idx += 1
            register_paragraph(block, f"s1.p{p_idx}")
            continue

        if block.__class__.__name__ != "Table":
            continue

        tbl_counter += 1
        p_idx += 1
        walk_table(block, f"s1.p{p_idx}.r1.tbl{tbl_counter}")

    return _EditableDocIndex(paragraphs=paragraphs, runs=runs, run_to_paragraph=run_to_paragraph)


_SECTION_NAME_RE = re.compile(r"^Contents/section(\d+)\.xml$")
_HP_NS = "http://www.hancom.co.kr/hwpml/2011/paragraph"
_HP = f"{{{_HP_NS}}}"


def _run_text(run_el: ET.Element) -> str:
    return "".join("".join(node.itertext()) for node in run_el.findall(f"{_HP}t"))


def _set_hwpx_run_text(run_el: ET.Element, new_text: str) -> None:
    for node in list(run_el):
        if node.tag == f"{_HP}t":
            run_el.remove(node)
    text_el = ET.SubElement(run_el, f"{_HP}t")
    text_el.text = new_text


def _iter_section_paragraphs(section_root: ET.Element) -> list[ET.Element]:
    return section_root.findall(f"{_HP}p")


def _iter_paragraph_tables(paragraph_el: ET.Element) -> list[ET.Element]:
    return paragraph_el.findall(f"{_HP}run/{_HP}tbl")


def _iter_cell_paragraphs(cell_el: ET.Element) -> list[ET.Element]:
    direct = cell_el.findall(f"{_HP}subList/{_HP}p")
    if direct:
        return direct
    return cell_el.findall(f".//{_HP}p")


def _safe_int(value: str | None) -> int | None:
    if value is None:
        return None
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def _logical_table_cells(row_el: ET.Element) -> list[tuple[int, ET.Element]]:
    logical_cells: list[tuple[int, ET.Element]] = []
    fallback_col = 1
    for cell_el in row_el.findall(f"{_HP}tc"):
        cell_addr = cell_el.find(f"{_HP}cellAddr")
        col_addr = _safe_int(cell_addr.get("colAddr")) if cell_addr is not None else None
        logical_col = (col_addr + 1) if col_addr is not None else fallback_col
        logical_cells.append((logical_col, cell_el))

        cell_span = cell_el.find(f"{_HP}cellSpan")
        colspan = _safe_int(cell_span.get("colSpan")) if cell_span is not None else None
        fallback_col = max(fallback_col, logical_col + max(colspan or 1, 1))
    return logical_cells


class _EditableHwpxArchive:
    def __init__(self, *, source_path: Path, source_bytes: bytes, section_entries: list[tuple[str, ET.Element]]) -> None:
        self.source_path = source_path
        self.source_bytes = source_bytes
        self.section_entries = section_entries

    @classmethod
    def open(cls, source_path: str | Path) -> "_EditableHwpxArchive":
        path = Path(source_path)
        source_bytes = path.read_bytes()
        with zipfile.ZipFile(BytesIO(source_bytes)) as archive:
            section_entries = sorted(
                [
                    (name, ET.fromstring(archive.read(name)))
                    for name in archive.namelist()
                    if _SECTION_NAME_RE.match(name)
                ],
                key=lambda item: int(_SECTION_NAME_RE.match(item[0]).group(1)),
            )
        return cls(source_path=path, source_bytes=source_bytes, section_entries=section_entries)

    def write_to(self, output_path: str | Path) -> None:
        ET.register_namespace("hs", "http://www.hancom.co.kr/hwpml/2011/section")
        ET.register_namespace("hp", _HP_NS)

        section_bytes = {
            name: ET.tostring(root, encoding="utf-8", xml_declaration=True)
            for name, root in self.section_entries
        }

        output = Path(output_path)
        with zipfile.ZipFile(BytesIO(self.source_bytes), "r") as source_archive:
            with zipfile.ZipFile(output, "w") as target_archive:
                for info in source_archive.infolist():
                    data = section_bytes.get(info.filename, source_archive.read(info.filename))
                    target_archive.writestr(info, data)


def _build_hwpx_index(archive: _EditableHwpxArchive) -> _EditableDocIndex:
    paragraphs: dict[str, _EditableParagraphRef] = {}
    runs: dict[str, _EditableRunRef] = {}
    run_to_paragraph: dict[str, _EditableParagraphRef] = {}

    def register_paragraph(paragraph_el: ET.Element, paragraph_id: str) -> _EditableParagraphRef:
        run_elements = paragraph_el.findall(f"{_HP}run")
        paragraph_ref = _EditableParagraphRef(
            unit_id=paragraph_id,
            runs=[],
            has_non_run_content=bool(_iter_paragraph_tables(paragraph_el)),
        )
        paragraphs[paragraph_id] = paragraph_ref
        if not run_elements:
            return paragraph_ref

        for run_index, run_el in enumerate(run_elements, start=1):
            run_id = f"{paragraph_id}.r{run_index}"
            run_ref = _EditableRunRef(
                unit_id=run_id,
                get_text=lambda node=run_el: _run_text(node),
                set_text=lambda value, node=run_el: _set_hwpx_run_text(node, value),
            )
            paragraph_ref.runs.append(run_ref)
            runs[run_id] = run_ref
            run_to_paragraph[run_id] = paragraph_ref
        return paragraph_ref

    def walk_table(table_el: ET.Element, table_base: str) -> None:
        for tr_idx, row_el in enumerate(table_el.findall(f"{_HP}tr"), start=1):
            for tc_idx, cell_el in _logical_table_cells(row_el):
                cell_paragraphs = _iter_cell_paragraphs(cell_el)
                if not cell_paragraphs:
                    paragraph_id = f"{table_base}.tr{tr_idx}.tc{tc_idx}.p1"
                    paragraphs[paragraph_id] = _EditableParagraphRef(
                        unit_id=paragraph_id,
                        runs=[],
                        has_non_run_content=False,
                    )
                    continue

                for cp_idx, paragraph_el in enumerate(cell_paragraphs, start=1):
                    paragraph_id = f"{table_base}.tr{tr_idx}.tc{tc_idx}.p{cp_idx}"
                    paragraph_ref = register_paragraph(paragraph_el, paragraph_id)
                    for nested_index, nested_table in enumerate(_iter_paragraph_tables(paragraph_el), start=1):
                        paragraph_ref.has_non_run_content = True
                        walk_table(nested_table, f"{paragraph_id}.tbl{nested_index}")

    for section_index, (_section_name, section_root) in enumerate(archive.section_entries, start=1):
        for paragraph_index, paragraph_el in enumerate(_iter_section_paragraphs(section_root), start=1):
            paragraph_id = f"s{section_index}.p{paragraph_index}"
            paragraph_ref = register_paragraph(paragraph_el, paragraph_id)
            for table_index, table_el in enumerate(_iter_paragraph_tables(paragraph_el), start=1):
                paragraph_ref.has_non_run_content = True
                walk_table(table_el, f"{paragraph_id}.r1.tbl{table_index}")

    return _EditableDocIndex(paragraphs=paragraphs, runs=runs, run_to_paragraph=run_to_paragraph)


def _build_run_spans(paragraph: _EditableParagraphRef) -> list[_RunSpan]:
    spans: list[_RunSpan] = []
    cursor = 0
    for run in paragraph.runs:
        length = len(run.text)
        spans.append(_RunSpan(start=cursor, end=cursor + length, run=run))
        cursor += length
    return spans


def _clip_run_spans(spans: list[_RunSpan], i1: int, i2: int) -> list[_RunSpan]:
    is_insert = i1 == i2
    clipped: list[_RunSpan] = []
    for span in spans:
        if is_insert:
            if span.end < i1 or span.start > i2:
                continue
        else:
            if span.end <= i1 or span.start >= i2:
                continue
        clipped.append(
            _RunSpan(
                start=span.start if is_insert else max(span.start, i1),
                end=span.end if is_insert else min(span.end, i2),
                run=span.run,
            )
        )
    if is_insert and len(clipped) > 1:
        preceding = [span for span in clipped if span.end == i1]
        if preceding:
            clipped = preceding[:1]
    return clipped


def _append_unique(values: list[str], value: str) -> None:
    if value not in values:
        values.append(value)


def _apply_to_run(run: _EditableRunRef, new_text: str, local_start: int, local_end: int) -> None:
    current = run.text
    run.text = current[:local_start] + new_text + current[local_end:]


def _apply_multi_run(
    orig_sub: str,
    new_sub: str,
    spans: list[_RunSpan],
    result: ApplyEditsResult,
    *,
    depth: int = 0,
) -> None:
    if depth == 0:
        matcher = difflib.SequenceMatcher(None, orig_sub, new_sub, autojunk=False)
        for tag, a1, a2, b1, b2 in matcher.get_opcodes():
            if tag == "equal":
                continue
            sub_spans = _clip_run_spans(spans, a1, a2)
            if not sub_spans:
                continue
            if len(sub_spans) == 1:
                span = sub_spans[0]
                local_start = a1 - span.start
                local_end = a2 - span.start
                _apply_to_run(span.run, new_sub[b1:b2], local_start, local_end)
                _append_unique(result.modified_run_ids, span.run.unit_id)
                continue
            _apply_multi_run(orig_sub[a1:a2], new_sub[b1:b2], sub_spans, result, depth=1)
        return

    first = spans[0].run
    first.text = new_sub
    _append_unique(result.modified_run_ids, first.unit_id)
    for span in spans[1:]:
        span.run.text = ""
        _append_unique(result.modified_run_ids, span.run.unit_id)
    result.warnings.append(
        "Multi-run fallback used for "
        f"{[span.run.unit_id for span in spans]}: all replacement text assigned to {first.unit_id}"
    )


def _validate_paragraph_edit(index: _EditableDocIndex, edit: ParagraphTextEdit) -> _EditableParagraphRef:
    paragraph = index.paragraphs.get(edit.paragraph_unit_id)
    if paragraph is None:
        raise EditValidationError(f"Paragraph does not exist: {edit.paragraph_unit_id}")
    if paragraph.has_non_run_content:
        raise EditValidationError(
            f"Paragraph edit targets unsupported mixed content (tables/images): {edit.paragraph_unit_id}"
        )
    if paragraph.text != edit.old_text:
        raise EditValidationError(
            f"Paragraph text mismatch for {edit.paragraph_unit_id}: expected {edit.old_text!r}, got {paragraph.text!r}"
        )
    return paragraph


def _validate_run_edit(index: _EditableDocIndex, edit: RunTextEdit) -> _EditableRunRef:
    run = index.runs.get(edit.run_unit_id)
    if run is None:
        raise EditValidationError(f"Run does not exist: {edit.run_unit_id}")
    if run.text != edit.old_text:
        raise EditValidationError(
            f"Run text mismatch for {edit.run_unit_id}: expected {edit.old_text!r}, got {run.text!r}"
        )
    return run


def _apply_single_edit(index: _EditableDocIndex, edit: EditCommand, result: ApplyEditsResult) -> None:
    if isinstance(edit, RunTextEdit):
        run = _validate_run_edit(index, edit)
        run.text = edit.new_text
        paragraph = index.run_to_paragraph.get(edit.run_unit_id)
        if paragraph is not None:
            paragraph.recompute()
        _append_unique(result.modified_unit_ids, edit.run_unit_id)
        _append_unique(result.modified_run_ids, edit.run_unit_id)
        result.edits_applied += 1
        return

    paragraph = _validate_paragraph_edit(index, edit)
    spans = _build_run_spans(paragraph)
    original = paragraph.text
    matcher = difflib.SequenceMatcher(None, original, edit.new_text, autojunk=False)
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            continue
        affected = _clip_run_spans(spans, i1, i2)
        if not affected:
            continue
        if len(affected) == 1:
            span = affected[0]
            local_start = i1 - span.start
            local_end = i2 - span.start
            _apply_to_run(span.run, edit.new_text[j1:j2], local_start, local_end)
            _append_unique(result.modified_run_ids, span.run.unit_id)
            continue
        _apply_multi_run(original[i1:i2], edit.new_text[j1:j2], affected, result)
    paragraph.recompute()
    _append_unique(result.modified_unit_ids, edit.paragraph_unit_id)
    result.edits_applied += 1


def validate_edit_commands(doc: DocIR, edits: list[EditCommand]) -> None:
    index = _build_doc_ir_index(doc)
    for edit in edits:
        if isinstance(edit, RunTextEdit):
            _validate_run_edit(index, edit)
        else:
            _validate_paragraph_edit(index, edit)


def apply_edits_to_doc_ir(doc: DocIR, edits: list[EditCommand]) -> tuple[DocIR, ApplyEditsResult]:
    updated = doc.model_copy(deep=True)
    index = _build_doc_ir_index(updated)
    result = ApplyEditsResult(source_doc_type=updated.source_doc_type)
    for edit in edits:
        _apply_single_edit(index, edit, result)
    return updated, result


def _default_output_path(source_path: Path) -> Path:
    return source_path.with_stem(f"{source_path.stem}_edited")


def apply_edits_to_file(
    source_path: str | Path,
    edits: list[EditCommand],
    *,
    output_path: str | Path | None = None,
) -> ApplyEditsResult:
    source = Path(source_path)
    doc = DocIR.from_file(source)
    result = ApplyEditsResult(source_doc_type=doc.source_doc_type)
    target_path = Path(output_path) if output_path is not None else _default_output_path(source)

    if doc.source_doc_type == "docx":
        from docx import Document as load_docx

        native_doc = load_docx(str(source))
        index = _build_docx_index(native_doc)
        for edit in edits:
            _apply_single_edit(index, edit, result)
        native_doc.save(str(target_path))
    elif doc.source_doc_type == "hwpx":
        archive = _EditableHwpxArchive.open(source)
        index = _build_hwpx_index(archive)
        for edit in edits:
            _apply_single_edit(index, edit, result)
        archive.write_to(target_path)
    else:
        raise EditValidationError(
            f"Native write-back is currently supported only for docx/hwpx, got {doc.source_doc_type!r}."
        )

    result.output_path = str(target_path)
    return result


__all__ = [
    "ApplyEditsResult",
    "EditValidationError",
    "ParagraphTextEdit",
    "RunTextEdit",
    "apply_edits_to_doc_ir",
    "apply_edits_to_file",
    "validate_edit_commands",
]
