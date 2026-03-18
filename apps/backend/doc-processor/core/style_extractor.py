"""
style_extractor — extract style information from HWPX / DOCX documents
=======================================================================

Produces a :class:`StyleMap` keyed by the same chunk IDs that the IR
pipeline uses, so styles can be recombined at render time without
polluting the LLM-facing IR.

Usage::

    from core.style_extractor import extract_styles_hwpx, extract_styles_docx

    style_map = extract_styles_hwpx(doc)   # HwpxDocument
    style_map = extract_styles_docx(path)  # Path to .docx
"""

from __future__ import annotations

import re
from pathlib import Path
from xml.etree import ElementTree as ET

from las_types import RunStyleInfo, ParaStyleInfo, CellStyleInfo, TableStyleInfo, StyleMap


# ---------------------------------------------------------------------------
# HWPX
# ---------------------------------------------------------------------------

_NS_HH = "http://www.hancom.co.kr/hwpml/2011/head"
_NS_HC = "http://www.hancom.co.kr/hwpml/2011/core"
_NS_HP = "http://www.hancom.co.kr/hwpml/2011/paragraph"

# HWPX border type → CSS border-style keyword
_HWPX_BORDER_STYLE = {
    "SOLID": "solid",
    "DASH": "dashed",
    "DOT": "dotted",
    "DASH_DOT": "dashed",
    "DASH_DOT_DOT": "dotted",
    "DOUBLE": "double",
    "NONE": "none",
}

# HWPX horizontal alignment → CSS text-align
_HWPX_HALIGN = {
    "LEFT": "left",
    "CENTER": "center",
    "RIGHT": "right",
    "JUSTIFY": "justify",
    "DISTRIBUTE": "justify",
}

# HWPX vertical alignment → CSS vertical-align
_HWPX_VALIGN = {
    "TOP": "top",
    "CENTER": "center",
    "BOTTOM": "bottom",
    "BASELINE": "top",
}


def _hwpx_border_css(border_el: ET.Element | None) -> str | None:
    """Convert a HWPX border element to CSS shorthand like ``'1px solid #000'``."""
    if border_el is None:
        return None
    btype = border_el.get("type", "NONE")
    if btype == "NONE":
        return None
    width = border_el.get("width", "0.12 mm")
    # Convert mm string to approximate px (1mm ≈ 3.78px, but 1px minimum is fine)
    try:
        mm_val = float(width.replace("mm", "").strip())
        px = max(1, round(mm_val * 3.78))
    except (ValueError, AttributeError):
        px = 1
    color = border_el.get("color", "#000000")
    style = _HWPX_BORDER_STYLE.get(btype, "solid")
    return f"{px}px {style} {color}"


def _build_border_fill_map(header) -> dict[str, ET.Element]:
    """Build borderFillIDRef → Element lookup from header XML."""
    mapping: dict[str, ET.Element] = {}
    for bf in header.element.findall(f".//{{{_NS_HH}}}borderFill"):
        bf_id = bf.get("id")
        if bf_id:
            mapping[bf_id] = bf
    return mapping


def _build_para_pr_map(header) -> dict[str, ET.Element]:
    """Build paraPrIDRef → Element lookup from header XML."""
    mapping: dict[str, ET.Element] = {}
    for pp in header.element.findall(f".//{{{_NS_HH}}}paraPr"):
        pp_id = pp.get("id")
        if pp_id:
            mapping[pp_id] = pp
    return mapping


def _hwpx_run_style(run) -> RunStyleInfo:
    """Extract RunStyleInfo from an HwpxOxmlRun."""
    style = run.style
    info = RunStyleInfo(
        bold=bool(run.bold),
        italic=bool(run.italic),
        underline=bool(run.underline),
    )

    if style is None:
        return info

    # Text color
    tc = style.text_color()
    if tc and tc != "#000000":
        info.color = tc

    # Font size: attributes["height"] is in 1/100 pt (HWP units)
    height = style.attributes.get("height")
    if height:
        try:
            info.size_pt = int(height) / 100.0
        except (ValueError, TypeError):
            pass

    # Strikethrough
    strikeout = style.child_attributes.get("strikeout", {})
    if strikeout.get("shape", "NONE") != "NONE":
        info.strikethrough = True

    # Underline details (already have bool, but check type too)
    ul = style.child_attributes.get("underline", {})
    if ul.get("type", "NONE") != "NONE":
        info.underline = True

    return info


def _hwpx_cell_style(
    cell,
    bf_map: dict[str, ET.Element],
    pp_map: dict[str, ET.Element],
) -> CellStyleInfo:
    """Extract CellStyleInfo from an HwpxOxmlTableCell."""
    info = CellStyleInfo()

    # Span (merge)
    span = cell.span  # (rowspan, colspan)
    info.rowspan = span[0]
    info.colspan = span[1]

    # Vertical alignment from subList
    sub = cell.element.find(f"{{{_NS_HP}}}subList")
    if sub is not None:
        valign = sub.get("vertAlign", "")
        info.vertical_align = _HWPX_VALIGN.get(valign)

    # Horizontal alignment from first paragraph's paraPrIDRef
    for cp in cell.paragraphs:
        pp_ref = cp.element.get("paraPrIDRef")
        if pp_ref and pp_ref in pp_map:
            align_el = pp_map[pp_ref].find(f"{{{_NS_HH}}}align")
            if align_el is not None:
                halign = align_el.get("horizontal", "")
                info.horizontal_align = _HWPX_HALIGN.get(halign)
        break

    # Borders and background from borderFillIDRef
    bf_ref = cell.element.get("borderFillIDRef")
    if bf_ref and bf_ref in bf_map:
        bf = bf_map[bf_ref]
        info.border_top = _hwpx_border_css(bf.find(f"{{{_NS_HH}}}topBorder"))
        info.border_bottom = _hwpx_border_css(bf.find(f"{{{_NS_HH}}}bottomBorder"))
        info.border_left = _hwpx_border_css(bf.find(f"{{{_NS_HH}}}leftBorder"))
        info.border_right = _hwpx_border_css(bf.find(f"{{{_NS_HH}}}rightBorder"))

        # Background fill — can be under hh: or hc: namespace
        fill_brush = bf.find(f"{{{_NS_HH}}}fillBrush") or bf.find(f"{{{_NS_HC}}}fillBrush")
        if fill_brush is not None:
            # faceColor can be on fillBrush itself or on a child winBrush
            face_color = fill_brush.get("faceColor")
            if not face_color:
                win_brush = (
                    fill_brush.find(f"{{{_NS_HH}}}winBrush")
                    or fill_brush.find(f"{{{_NS_HC}}}winBrush")
                )
                if win_brush is not None:
                    face_color = win_brush.get("faceColor")
            if face_color and face_color.lower() not in ("none", "#ffffff", "transparent"):
                info.background = face_color

    return info


def extract_styles_hwpx(doc) -> StyleMap:
    """Extract all style info from an open :class:`HwpxDocument`.

    Args:
        doc: An open ``HwpxDocument`` instance.

    Returns:
        :class:`StyleMap` keyed by the same chunk IDs as the IR.
    """
    sm = StyleMap()
    header = doc.headers[0] if doc.headers else None
    bf_map = _build_border_fill_map(header) if header else {}
    pp_map = _build_para_pr_map(header) if header else {}

    for s_idx, section in enumerate(doc.sections):
        for p_idx, para in enumerate(section.paragraphs):
            base = f"s{s_idx + 1}.p{p_idx + 1}"

            # Paragraph alignment
            pp_ref = para.element.get("paraPrIDRef")
            if pp_ref and pp_ref in pp_map:
                align_el = pp_map[pp_ref].find(f"{{{_NS_HH}}}align")
                if align_el is not None:
                    halign = _HWPX_HALIGN.get(align_el.get("horizontal", ""))
                    if halign:
                        sm.paragraphs[base] = ParaStyleInfo(align=halign)

            # Body runs
            for r_idx, run in enumerate(para.runs):
                chunk_id = f"{base}.r{r_idx + 1}"
                sm.runs[chunk_id] = _hwpx_run_style(run)

            # Tables — HWPX nests each table in its own paragraph,
            # so the exporter always numbers them tbl1 within that paragraph.
            for tbl_idx, tbl in enumerate(para.tables, start=1):
                tbl_root = f"{base}.r1.tbl{tbl_idx}"
                sm.tables[tbl_root] = TableStyleInfo(
                    row_count=tbl.row_count,
                    col_count=tbl.column_count,
                )

                cell_map = tbl.get_cell_map()
                seen_cells: set[str] = set()
                for row_positions in cell_map:
                    for pos in row_positions:
                        r, c = pos.anchor
                        cell_key = f"{tbl_root}.tr{r + 1}.tc{c + 1}"
                        if cell_key in seen_cells:
                            continue
                        seen_cells.add(cell_key)

                        sm.cells[cell_key] = _hwpx_cell_style(
                            pos.cell, bf_map, pp_map,
                        )

                        # Cell paragraph alignment and run styles
                        for cp_idx, cp in enumerate(pos.cell.paragraphs):
                            cp_pp_ref = cp.element.get("paraPrIDRef")
                            if cp_pp_ref and cp_pp_ref in pp_map:
                                cp_align_el = pp_map[cp_pp_ref].find(f"{{{_NS_HH}}}align")
                                if cp_align_el is not None:
                                    cp_halign = _HWPX_HALIGN.get(cp_align_el.get("horizontal", ""))
                                    if cp_halign:
                                        cp_key = f"{tbl_root}.tr{r + 1}.tc{c + 1}.p{cp_idx + 1}"
                                        sm.paragraphs[cp_key] = ParaStyleInfo(align=cp_halign)

                            for cr_idx, cr in enumerate(cp.runs):
                                run_id = f"{tbl_root}.tr{r + 1}.tc{c + 1}.p{cp_idx + 1}.r{cr_idx + 1}"
                                sm.runs[run_id] = _hwpx_run_style(cr)

    return sm


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _docx_run_style(run) -> RunStyleInfo:
    """Extract RunStyleInfo from a python-docx Run."""
    info = RunStyleInfo(
        bold=bool(run.bold),
        italic=bool(run.italic),
        underline=bool(run.underline),
    )

    font = run.font

    # Color
    if font.color and font.color.rgb:
        rgb = str(font.color.rgb)
        if rgb != "000000":
            info.color = f"#{rgb}"

    # Size
    if font.size is not None:
        info.size_pt = font.size / 12700  # EMU → pt

    # Strikethrough
    if font.strike:
        info.strikethrough = True

    # Super/subscript
    if font.superscript:
        info.superscript = True
    if font.subscript:
        info.subscript = True

    # Highlight
    if font.highlight_color is not None:
        info.highlight = str(font.highlight_color)

    return info


def _docx_border_css(tcBorders, side: str) -> str | None:
    """Extract CSS border from a w:tcBorders element."""
    if tcBorders is None:
        return None
    from docx.oxml.ns import qn
    el = tcBorders.find(qn(f"w:{side}"))
    if el is None:
        return None
    val = el.get(qn("w:val"), "none")
    if val in ("none", "nil"):
        return None
    sz = el.get(qn("w:sz"), "4")  # in 1/8 pt
    try:
        px = max(1, round(int(sz) / 8 * 1.333))
    except (ValueError, TypeError):
        px = 1
    color = el.get(qn("w:color"), "000000")
    style_map = {"single": "solid", "double": "double", "dashed": "dashed", "dotted": "dotted"}
    css_style = style_map.get(val, "solid")
    return f"{px}px {css_style} #{color}"


def _docx_cell_style(cell) -> CellStyleInfo:
    """Extract CellStyleInfo from a python-docx table Cell."""
    from docx.oxml.ns import qn

    info = CellStyleInfo()
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        return info

    # Horizontal merge (gridSpan)
    gs = tcPr.find(qn("w:gridSpan"))
    if gs is not None:
        try:
            info.colspan = int(gs.get(qn("w:val"), "1"))
        except (ValueError, TypeError):
            pass

    # Vertical alignment
    vAlign = tcPr.find(qn("w:vAlign"))
    if vAlign is not None:
        val = vAlign.get(qn("w:val"), "")
        info.vertical_align = {"top": "top", "center": "center", "bottom": "bottom"}.get(val)

    # Horizontal alignment from first paragraph
    for p in cell.paragraphs:
        if p.alignment is not None:
            align_map = {0: "left", 1: "center", 2: "right", 3: "justify"}
            info.horizontal_align = align_map.get(p.alignment, None)
        break

    # Shading / background
    shd = tcPr.find(qn("w:shd"))
    if shd is not None:
        fill = shd.get(qn("w:fill"))
        if fill and fill.lower() not in ("auto", "ffffff", "none"):
            info.background = f"#{fill}"

    # Borders
    tcBorders = tcPr.find(qn("w:tcBorders"))
    info.border_top = _docx_border_css(tcBorders, "top")
    info.border_bottom = _docx_border_css(tcBorders, "bottom")
    info.border_left = _docx_border_css(tcBorders, "left")
    info.border_right = _docx_border_css(tcBorders, "right")

    return info


def extract_styles_docx(source: str | Path) -> StyleMap:
    """Extract all style info from a DOCX file.

    Args:
        source: Path to the ``.docx`` file.

    Returns:
        :class:`StyleMap` keyed by the same chunk IDs as the IR.
    """
    from docx import Document as DocxDocument
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = DocxDocument(str(source))
    sm = StyleMap()

    p_idx = 0
    tbl_counter = 0

    _DOCX_ALIGN = {0: "left", 1: "center", 2: "right", 3: "justify"}

    for block in doc.iter_inner_content():
        if isinstance(block, Paragraph):
            p_idx += 1
            base = f"s1.p{p_idx}"

            # Paragraph alignment
            if block.alignment is not None:
                align = _DOCX_ALIGN.get(block.alignment)
                if align:
                    sm.paragraphs[base] = ParaStyleInfo(align=align)

            for r_idx, run in enumerate(block.runs, start=1):
                sm.runs[f"{base}.r{r_idx}"] = _docx_run_style(run)

        elif isinstance(block, Table):
            tbl_counter += 1
            p_idx += 1
            tbl_root = f"s1.p{p_idx}.r1.tbl{tbl_counter}"
            sm.tables[tbl_root] = TableStyleInfo(
                row_count=len(block.rows),
                col_count=len(block.columns),
            )

            # Track vertical merges to compute rowspan
            # vMerge with val="restart" starts a span; absent val continues it
            from docx.oxml.ns import qn
            vmerge_starts: dict[int, tuple[int, int]] = {}  # col → (start_row, start_col)

            for tr_idx, row in enumerate(block.rows, start=1):
                for tc_idx, cell in enumerate(row.cells, start=1):
                    cell_key = f"{tbl_root}.tr{tr_idx}.tc{tc_idx}"
                    style = _docx_cell_style(cell)

                    # Vertical merge detection
                    tc_el = cell._tc
                    tcPr = tc_el.find(qn("w:tcPr"))
                    if tcPr is not None:
                        vMerge = tcPr.find(qn("w:vMerge"))
                        if vMerge is not None:
                            val = vMerge.get(qn("w:val"), "")
                            if val == "restart":
                                vmerge_starts[tc_idx] = (tr_idx, tc_idx)
                            # continuation cells don't get their own entry
                            # (the start cell's rowspan will be incremented)
                            elif tc_idx in vmerge_starts:
                                start_r, start_c = vmerge_starts[tc_idx]
                                start_key = f"{tbl_root}.tr{start_r}.tc{start_c}"
                                if start_key in sm.cells:
                                    sm.cells[start_key].rowspan += 1
                                continue
                        else:
                            # No vMerge — end any active vertical span
                            vmerge_starts.pop(tc_idx, None)

                    sm.cells[cell_key] = style

                    # Cell paragraph alignment and run styles
                    for cp_idx, cp in enumerate(cell.paragraphs, start=1):
                        if cp.alignment is not None:
                            cp_align = _DOCX_ALIGN.get(cp.alignment)
                            if cp_align:
                                cp_key = f"{tbl_root}.tr{tr_idx}.tc{tc_idx}.p{cp_idx}"
                                sm.paragraphs[cp_key] = ParaStyleInfo(align=cp_align)
                        for cr_idx, cr in enumerate(cp.runs, start=1):
                            run_id = f"{tbl_root}.tr{tr_idx}.tc{tc_idx}.p{cp_idx}.r{cr_idx}"
                            sm.runs[run_id] = _docx_run_style(cr)

    return sm
