"""Generate DOCX sample contracts from markdown source files."""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "doc_samples/계약서_test_생성본(docx)")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def set_font(run, size_pt=11, bold=False, name="맑은 고딕"):
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size_pt=16, bold=True)
    p.paragraph_format.space_after = Pt(18)


def add_article_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size_pt=11, bold=True)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)


def add_body_line(doc: Document, text: str, indent: bool = False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run, size_pt=11)
    p.paragraph_format.space_after = Pt(2)


def add_date_line(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size_pt=11)
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)


def add_signature_line(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(2)
    run = p.add_run(text)
    set_font(run, size_pt=11)
    p.paragraph_format.space_after = Pt(4)


def setup_doc() -> Document:
    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(3.0)
    # Remove default paragraph spacing
    doc.styles["Normal"].paragraph_format.space_after = Pt(0)
    return doc


def parse_and_add_content(doc: Document, lines: list[str]):
    """Parse lines into document elements."""
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        # Title line (first line, starts with [L-XX])
        if line.startswith("[L-"):
            title = line[line.index("]") + 1:].strip()
            add_title(doc, title)
        # Article heading: 제N조(...)
        elif line.startswith("제") and "조" in line and line[1:].split("조")[0].replace("일", "1").isdigit() is False:
            # detect 제1조 style
            if "(" in line or line.endswith(")"):
                add_article_heading(doc, line)
            else:
                add_article_heading(doc, line)
        # Numbered items: 1. 2. 3. ...
        elif len(line) >= 2 and line[0].isdigit() and line[1] == ".":
            add_body_line(doc, line, indent=True)
        # Date line
        elif line.endswith("일") and "년" in line and "월" in line and len(line) < 20:
            add_date_line(doc, line)
        # Signature lines
        elif line.startswith("사용자:") or line.startswith("근로자:") or \
             line.startswith("회사:") or line.startswith("직원:") or \
             line.startswith("위탁자:") or line.startswith("수탁자:"):
            add_signature_line(doc, line)
        else:
            add_body_line(doc, line)
        i += 1


def make_article_heading_detector(line: str) -> bool:
    """Return True if line is an article heading like 제1조(...)."""
    if not line.startswith("제"):
        return False
    rest = line[1:]
    for i, ch in enumerate(rest):
        if ch == "조":
            num_part = rest[:i]
            # Korean number chars
            korean_nums = set("일이삼사오육칠팔구십")
            if all(c in korean_nums for c in num_part):
                return True
    return False


def build_docx(title_slug: str, content: str):
    doc = setup_doc()
    lines = content.splitlines()

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        # Title: [L-XX] ...
        if line.startswith("[L-"):
            title_text = line[line.index("]") + 1:].strip()
            add_title(doc, title_text)

        # Article heading
        elif make_article_heading_detector(line):
            add_article_heading(doc, line)

        # Numbered items
        elif len(line) >= 2 and line[0].isdigit() and line[1] == ".":
            add_body_line(doc, line, indent=True)

        # Date line: e.g. "2026년 3월 28일"
        elif "년" in line and "월" in line and line.rstrip().endswith("일") and len(line) < 25:
            add_date_line(doc, line)

        # Signature lines
        elif any(line.startswith(prefix) for prefix in
                 ("사용자:", "근로자:", "회사:", "직원:", "위탁자:", "수탁자:")):
            add_signature_line(doc, line)

        else:
            add_body_line(doc, line)

        i += 1

    out_path = os.path.join(OUTPUT_DIR, f"{title_slug}.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")


CONTRACTS = {
    "정규직_근로계약서": """\
[L-01] 정규직 근로계약서

제1조(당사자)
사용자: 주식회사 새온커머스(이하 "회사"라 한다)
근로자: 김민수

제2조(입사일)
근로자의 입사일은 2026년 4월 1일로 한다.

제3조(근무장소 및 업무)
1. 근무장소: 서울특별시 구로구 디지털로 214, 7층
2. 담당업무: 고객응대, 주문관리, 반품접수 및 이에 부수하는 업무

제4조(근로형태)
본 계약은 기간의 정함이 없는 근로계약으로 한다.

제5조(근로시간 및 휴게)
1. 소정근로시간은 월요일부터 금요일까지 09:00부터 18:00까지로 한다.
2. 휴게시간은 12:00부터 13:00까지로 한다.

제6조(휴일 및 휴가)
1. 주휴일은 일요일로 한다.
2. 근로자의 날 및 법령상 유급휴일은 유급으로 한다.
3. 연차유급휴가는 관계 법령 및 회사 취업규칙에 따른다.

제7조(임금)
1. 기본급: 월 2,550,000원
2. 식대: 월 100,000원
3. 연장·야간·휴일근로가 발생한 경우 관계 법령에 따라 별도 지급한다.
4. 임금지급일은 매월 10일로 하며, 전월 1일부터 말일까지의 임금을 지급한다.

제8조(기타)
1. 회사는 근로자에게 4대보험을 관계 법령에 따라 적용한다.
2. 근로자의 고의 또는 과실로 인한 재고손실, 고객환불, 민원보상금 등이 발생한 경우 회사는 그 금액을 산정하여 해당 월 임금에서 공제할 수 있다.

제9조(계약서 교부)
회사는 본 계약서 1부를 작성하여 근로자에게 교부한다.

2026년 3월 28일

사용자: 주식회사 새온커머스 대표이사 박성호 (인)
근로자: 김민수 (인)
""",
    "기간제_근로계약서": """\
[L-02] 기간제 근로계약서(근로조건통지서 겸용)

제1조(당사자)
사용자: 주식회사 도시랩
근로자: 이서현

제2조(계약기간)
2026년 5월 1일부터 2027년 4월 30일까지

제3조(근무장소 및 업무)
1. 근무장소: 서울특별시 마포구 월드컵북로 402, 9층
2. 담당업무: 데이터 정리, 문서 스캔, 자료 입력, 팀 지원업무

제4조(근로시간)
1. 근로일: 월요일부터 금요일까지
2. 근로시간: 10:00부터 17:00까지
3. 휴게시간: 13:00부터 14:00까지

제5조(임금)
1. 시급: 11,800원
2. 교통보조비: 월 100,000원
3. 임금지급일: 매월 15일
4. 지급방법: 근로자 명의 계좌로 입금

제6조(기타)
1. 계약기간 만료 시 본 계약은 별도의 통보 없이 종료될 수 있다.
2. 기타 정하지 않은 사항은 회사 취업규칙에 따른다.

2026년 4월 28일

사용자: 주식회사 도시랩 대표이사 한지우 (인)
근로자: 이서현 (인)
""",
    "연봉계약서": """\
[L-03] 연봉계약서

제1조(당사자)
회사: 주식회사 하늘데이터
직원: 박지윤

제2조(적용기간)
본 연봉계약의 적용기간은 2026년 4월 1일부터 2027년 3월 31일까지로 한다.

제3조(연봉)
1. 연봉 총액: 38,400,000원
2. 월 지급액: 3,200,000원
3. 임금지급일: 매월 25일
4. 위 금액에는 기본급, 직무수당 및 식대가 포함된다.

제4조(퇴직금)
상기 연봉총액에는 법정퇴직금이 포함되어 있으며, 회사는 연봉총액의 1/12를 매월 지급함으로써 퇴직금 지급을 갈음한다.

제5조(기타)
1. 그 밖의 근로조건은 기존 근로계약서와 취업규칙을 따른다.
2. 본 계약은 연봉조건에 관한 합의서로서 2부 작성하여 각 1부씩 보관한다.

2026년 3월 31일

회사: 주식회사 하늘데이터 대표이사 정유진 (인)
직원: 박지윤 (인)
""",
    "수습_근로계약서": """\
[L-04] 수습 근로계약서

제1조(당사자)
사용자: 주식회사 바른리테일
근로자: 최도윤

제2조(계약기간)
본 계약은 2026년 4월 1일부터 2027년 2월 28일까지로 한다.

제3조(근무장소 및 업무)
1. 근무장소: 경기도 성남시 분당구 판교역로 221
2. 담당업무: 온라인 주문 검수, 출고 확인, 재고 입력

제4조(근로시간 및 휴게)
1. 근로시간은 월요일부터 금요일까지 09:00부터 18:00까지로 한다.
2. 휴게시간은 12:00부터 13:00까지로 한다.

제5조(휴일 및 휴가)
1. 주휴일은 일요일로 한다.
2. 연차유급휴가는 관계 법령에 따른다.

제6조(임금)
1. 수습기간은 2026년 4월 1일부터 2026년 6월 30일까지 3개월로 한다.
2. 수습기간 중 임금은 월 1,980,000원으로 한다.
3. 수습기간 종료 후 임금은 월 2,200,000원으로 한다.
4. 임금지급일은 매월 10일로 한다.

제7조(기타)
본 계약서에 정하지 아니한 사항은 관계 법령 및 취업규칙에 따른다.

2026년 3월 30일

사용자: 주식회사 바른리테일 대표이사 오준석 (인)
근로자: 최도윤 (인)
""",
    "업무위탁계약서": """\
[L-05] 업무위탁계약서

제1조(당사자)
위탁자: 주식회사 라이트링크(이하 "회사"라 한다)
수탁자: 최서윤

제2조(계약목적)
수탁자는 회사가 운영하는 온라인 서비스의 게시물 검수 및 고객문의 1차 응대 업무를 수행한다.

제3조(계약기간)
2026년 4월 1일부터 2026년 9월 30일까지

제4조(업무수행 방식)
1. 수탁자는 회사 사무실(서울특별시 강남구 테헤란로 144)에서 월요일부터 금요일까지 09:00부터 18:00까지 근무한다.
2. 수탁자는 회사 팀장이 배정하는 일일 업무목록에 따라 업무를 수행한다.
3. 수탁자는 지각, 조퇴, 결근 또는 휴무가 필요한 경우 사전에 팀장 승인을 받아야 한다.
4. 수탁자는 회사가 제공하는 PC, 사내 메신저 및 업무계정을 사용한다.
5. 수탁자는 매일 업무종료 전 일일업무보고서를 제출한다.

제5조(보수)
1. 업무위탁 수수료는 월 2,700,000원으로 한다.
2. 지급일은 매월 10일로 한다.

제6조(법적 성격)
본 계약은 도급 또는 위임계약이며, 수탁자에게는 근로기준법상 연차휴가, 연장근로수당, 퇴직금 및 4대보험을 적용하지 않는다.

제7조(기타)
본 계약에서 정하지 않은 사항은 상호 협의하여 정한다.

2026년 3월 29일

위탁자: 주식회사 라이트링크 대표이사 임하린 (인)
수탁자: 최서윤 (인)
""",
}


if __name__ == "__main__":
    for slug, content in CONTRACTS.items():
        build_docx(slug, content)
    print("Done.")
